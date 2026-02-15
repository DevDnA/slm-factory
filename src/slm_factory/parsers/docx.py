"""DOCX (Word) 문서 파서입니다."""

from __future__ import annotations

from pathlib import Path
from typing import ClassVar

from ..models import ParsedDocument
from ..utils import get_logger
from .base import BaseParser, extract_date_from_filename

logger = get_logger("parsers.docx")


class DOCXParser(BaseParser):
    """DOCX 문서를 파싱합니다."""

    extensions: ClassVar[list[str]] = [".docx"]

    def parse(self, path: Path) -> ParsedDocument:
        """DOCX 파일을 파싱하여 ParsedDocument를 반환합니다.

        매개변수
        ----------
        path:
            DOCX 파일의 경로입니다.

        반환값
        -------
        ParsedDocument
            텍스트 내용, 표 (마크다운), 메타데이터로 채워집니다.

        예외
        ------
        ImportError
            python-docx가 설치되지 않았으면 발생합니다.
        FileNotFoundError
            *path*가 존재하지 않으면 발생합니다.
        RuntimeError
            DOCX 파일을 파싱할 수 없으면 발생합니다.
        """
        try:
            from docx import Document
        except ImportError as exc:
            raise ImportError(
                "python-docx가 설치되지 않았습니다. "
                "pip install slm-factory[docx] 로 설치하세요."
            ) from exc

        path = Path(path)
        if not path.exists():
            raise FileNotFoundError(f"DOCX file not found: {path}")

        # ------------------------------------------------------------------
        # DOCX 파일 로드
        # ------------------------------------------------------------------
        try:
            doc = Document(str(path))
        except Exception as exc:
            raise RuntimeError(f"Failed to parse DOCX: {path}") from exc

        # ------------------------------------------------------------------
        # 텍스트 추출
        # ------------------------------------------------------------------
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # 헤딩 스타일 처리
                if para.style and hasattr(para.style, "name") and isinstance(para.style.name, str):
                    if para.style.name.startswith("Heading"):
                        # Heading 1, Heading 2 등에서 숫자 추출
                        level = 1
                        if para.style.name[-1].isdigit():
                            level = int(para.style.name[-1])
                        paragraphs.append(f"{'#' * level} {text}")
                    else:
                        paragraphs.append(text)
                else:
                    paragraphs.append(text)

        content = "\n\n".join(paragraphs)

        # ------------------------------------------------------------------
        # 표 추출
        # ------------------------------------------------------------------
        tables_md: list[str] = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)

            if rows:
                # 마크다운 표로 변환
                header = "| " + " | ".join(rows[0]) + " |"
                separator = "| " + " | ".join(["---"] * len(rows[0])) + " |"

                if len(rows) > 1:
                    body_lines = []
                    for row in rows[1:]:
                        # 행을 헤더 길이에 맞게 패딩
                        while len(row) < len(rows[0]):
                            row.append("")
                        body_lines.append("| " + " | ".join(row[: len(rows[0])]) + " |")
                    body = "\n".join(body_lines)
                    md_table = f"{header}\n{separator}\n{body}"
                else:
                    md_table = f"{header}\n{separator}"

                tables_md.append(md_table)

        # ------------------------------------------------------------------
        # 제목 추출
        # ------------------------------------------------------------------
        title = path.stem
        if doc.core_properties.title:
            title = doc.core_properties.title

        # ------------------------------------------------------------------
        # 메타데이터
        # ------------------------------------------------------------------
        metadata: dict[str, object] = {}
        props = doc.core_properties

        if props.author:
            metadata["author"] = props.author
        if props.created:
            metadata["date"] = props.created.strftime("%Y-%m-%d")

        # 파일명에서 날짜 추출 시도 (YYMMDD)
        date_from_name = extract_date_from_filename(path.stem)
        if date_from_name and "date" not in metadata:
            metadata["date"] = date_from_name

        return ParsedDocument(
            doc_id=path.name,
            title=title,
            content=content,
            tables=tables_md,
            metadata=metadata,
        )
